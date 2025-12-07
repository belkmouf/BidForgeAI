"""
Document Processing - Extract text from various file formats
"""

import io
from typing import Dict, List, Optional
import PyPDF2
import pdfplumber
from pathlib import Path


class DocumentProcessor:
    """Process and extract text from various document formats"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.docx', '.xlsx', '.msg']

    def process_file(self, file_bytes: bytes, filename: str) -> Dict[str, any]:
        """
        Process uploaded file and extract text content

        Args:
            file_bytes: Raw file bytes
            filename: Original filename

        Returns:
            Dict with extracted text, metadata, and processing status
        """

        file_ext = Path(filename).suffix.lower()

        try:
            if file_ext == '.pdf':
                return self.process_pdf(file_bytes, filename)
            elif file_ext == '.txt':
                return self.process_txt(file_bytes, filename)
            elif file_ext == '.docx':
                return self.process_docx(file_bytes, filename)
            elif file_ext == '.xlsx':
                return self.process_xlsx(file_bytes, filename)
            elif file_ext == '.msg':
                return self.process_msg(file_bytes, filename)
            else:
                return {
                    'success': False,
                    'error': f'Unsupported file format: {file_ext}',
                    'filename': filename
                }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'filename': filename
            }

    def process_pdf(self, file_bytes: bytes, filename: str) -> Dict:
        """Extract text from PDF file"""

        text_content = []
        metadata = {}

        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                metadata['pages'] = len(pdf.pages)
                metadata['format'] = 'PDF'

                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_content.append(f"--- Page {page_num} ---\n{text}")

                    # Extract tables if any
                    tables = page.extract_tables()
                    if tables:
                        for table_num, table in enumerate(tables, 1):
                            text_content.append(f"\n[Table {table_num} on Page {page_num}]")
                            # Convert table to text representation
                            for row in table:
                                text_content.append(" | ".join(str(cell) if cell else "" for cell in row))

            full_text = "\n\n".join(text_content)

            return {
                'success': True,
                'filename': filename,
                'content': full_text,
                'metadata': metadata,
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }

        except Exception as e:
            # Fallback to PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                metadata['pages'] = len(pdf_reader.pages)
                metadata['format'] = 'PDF'

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_content.append(f"--- Page {page_num} ---\n{text}")

                full_text = "\n\n".join(text_content)

                return {
                    'success': True,
                    'filename': filename,
                    'content': full_text,
                    'metadata': metadata,
                    'word_count': len(full_text.split()),
                    'char_count': len(full_text)
                }
            except Exception as fallback_error:
                return {
                    'success': False,
                    'error': f"PDF processing failed: {str(e)}, Fallback: {str(fallback_error)}",
                    'filename': filename
                }

    def process_txt(self, file_bytes: bytes, filename: str) -> Dict:
        """Extract text from TXT file"""

        try:
            # Try UTF-8 first
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            text = file_bytes.decode('latin-1')

        return {
            'success': True,
            'filename': filename,
            'content': text,
            'metadata': {'format': 'TXT'},
            'word_count': len(text.split()),
            'char_count': len(text)
        }

    def process_docx(self, file_bytes: bytes, filename: str) -> Dict:
        """Extract text from DOCX file"""

        try:
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    tables_text.append(row_text)

            full_text = "\n\n".join(paragraphs)
            if tables_text:
                full_text += "\n\n[TABLES]\n" + "\n".join(tables_text)

            return {
                'success': True,
                'filename': filename,
                'content': full_text,
                'metadata': {
                    'format': 'DOCX',
                    'paragraphs': len(paragraphs),
                    'tables': len(doc.tables)
                },
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }

        except Exception as e:
            return {
                'success': False,
                'error': f"DOCX processing failed: {str(e)}",
                'filename': filename
            }

    def process_xlsx(self, file_bytes: bytes, filename: str) -> Dict:
        """Extract text from Excel file"""

        try:
            import pandas as pd

            # Read all sheets
            excel_file = pd.ExcelFile(io.BytesIO(file_bytes))
            text_content = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text_content.append(f"=== Sheet: {sheet_name} ===\n")
                text_content.append(df.to_string())

            full_text = "\n\n".join(text_content)

            return {
                'success': True,
                'filename': filename,
                'content': full_text,
                'metadata': {
                    'format': 'XLSX',
                    'sheets': len(excel_file.sheet_names),
                    'sheet_names': excel_file.sheet_names
                },
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }

        except Exception as e:
            return {
                'success': False,
                'error': f"XLSX processing failed: {str(e)}",
                'filename': filename
            }

    def process_msg(self, file_bytes: bytes, filename: str) -> Dict:
        """Extract text from Outlook MSG file"""

        try:
            import extract_msg

            msg = extract_msg.Message(io.BytesIO(file_bytes))

            content_parts = []
            content_parts.append(f"From: {msg.sender}")
            content_parts.append(f"To: {msg.to}")
            content_parts.append(f"Subject: {msg.subject}")
            content_parts.append(f"Date: {msg.date}")
            content_parts.append(f"\n{msg.body}")

            # Extract attachments info
            if msg.attachments:
                content_parts.append(f"\n[Attachments: {len(msg.attachments)}]")
                for attachment in msg.attachments:
                    content_parts.append(f"  - {attachment.longFilename}")

            full_text = "\n".join(content_parts)

            return {
                'success': True,
                'filename': filename,
                'content': full_text,
                'metadata': {
                    'format': 'MSG',
                    'from': msg.sender,
                    'subject': msg.subject,
                    'date': str(msg.date),
                    'attachments': len(msg.attachments) if msg.attachments else 0
                },
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }

        except Exception as e:
            return {
                'success': False,
                'error': f"MSG processing failed: {str(e)}",
                'filename': filename
            }

    def chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        overlap: int = 200
    ) -> List[str]:
        """
        Split text into overlapping chunks for RAG processing

        Args:
            text: Text to chunk
            chunk_size: Maximum chunk size in characters
            overlap: Overlap between chunks

        Returns:
            List of text chunks
        """

        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                for sep in ['. ', '.\n', '! ', '?\n']:
                    last_sep = text[start:end].rfind(sep)
                    if last_sep > chunk_size * 0.5:  # At least 50% of chunk size
                        end = start + last_sep + len(sep)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap if end < len(text) else end

        return chunks

    def extract_key_info(self, text: str) -> Dict:
        """
        Extract key information from document text
        (Project names, dates, budget figures, etc.)
        """

        import re

        info = {}

        # Extract dates (various formats)
        date_patterns = [
            r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',
            r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',
            r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}'
        ]

        dates = []
        for pattern in date_patterns:
            dates.extend(re.findall(pattern, text, re.IGNORECASE))
        info['dates'] = list(set(dates))[:10]  # Limit to 10 unique dates

        # Extract monetary values
        money_pattern = r'\$\s*\d+(?:,\d{3})*(?:\.\d{2})?(?:\s*(?:million|M|billion|B|thousand|K))?'
        info['monetary_values'] = re.findall(money_pattern, text, re.IGNORECASE)[:20]

        # Extract percentages
        percentage_pattern = r'\d+(?:\.\d+)?%'
        info['percentages'] = re.findall(percentage_pattern, text)[:20]

        # Extract email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        info['emails'] = re.findall(email_pattern, text)[:10]

        # Extract phone numbers
        phone_pattern = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}'
        info['phone_numbers'] = re.findall(phone_pattern, text)[:10]

        return info


# Global document processor instance
doc_processor = DocumentProcessor()
